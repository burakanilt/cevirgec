import os
import fitz
from PIL import Image, ImageDraw, ImageFont
from core.etds_tools import generate_ek1_decision, generate_ek2_user_notice, extract_verification_code
from core.pdf_ops import apply_bottom_margin
from docx import Document

def run_tests():
    print("--- SPRINT 5 ETDS KANIT ZİNCİRİ TESTLERİ ---")
    
    # 1. EK-1
    data_ek1 = {
        "sirket_unvani": "Test A.Ş.",
        "karar_tarihi": "09.07.2026",
        "karar_no": "2026/05",
        "katilanlar": "Ahmet Yılmaz, Ayşe Demir",
        "yetkili_ad_soyad": "Ahmet Yılmaz",
        "yetkili_tckn": "12345678901",
        "yetkili_eposta": "ahmet@test.com",
        "yetkili_telefon": "05551234567",
        "defter_turleri": "Yönetim Kurulu Karar Defteri"
    }
    ek1_path = "tests/test_ek1.docx"
    generate_ek1_decision(data_ek1, ek1_path)
    doc_ek1 = Document(ek1_path)
    full_text_ek1 = "\n".join([p.text for p in doc_ek1.paragraphs])
    print(f"[1] EK-1 üretildi: {ek1_path}")
    print(f"    İçerik kontrolü (TCKN): {'12345678901' in full_text_ek1}")
    
    # 2. EK-2
    data_ek2 = {
        "sirket_unvani": "Test A.Ş.",
        "bildirim_tarihi": "10.07.2026",
        "yeni_kullanicilar": [{"ad_soyad": "Veli Çelik", "tckn": "98765432101", "eposta": "veli@test.com", "telefon": "05559876543"}],
        "kaldırilan_kullanicilar": [{"ad_soyad": "Ayşe Demir", "tckn": "11122233344"}]
    }
    ek2_path = "tests/test_ek2.docx"
    generate_ek2_user_notice(data_ek2, ek2_path)
    doc_ek2 = Document(ek2_path)
    full_text_ek2 = " ".join([cell.text for table in doc_ek2.tables for row in table.rows for cell in row.cells])
    print(f"[2] EK-2 üretildi: {ek2_path}")
    print(f"    Yeni kullanıcı (Veli Çelik) tabloda mı?: {'Veli Çelik' in full_text_ek2}")
    print(f"    Kaldırılan kullanıcı (Ayşe Demir) tabloda mı?: {'Ayşe Demir' in full_text_ek2}")

    # 3. Marj Uyumlulaştırıcı
    # Create a dummy PDF first
    dummy_pdf_path = "tests/dummy_karar.pdf"
    doc_dummy = fitz.open()
    page = doc_dummy.new_page(width=595, height=842) # A4
    page.insert_text(fitz.Point(50, 800), "Alt Kısıma Çok Yakın Yazı!") # Bottom is 842
    doc_dummy.save(dummy_pdf_path)
    doc_dummy.close()
    
    out_pdf_path = "tests/dummy_karar_margin.pdf"
    apply_bottom_margin(dummy_pdf_path, out_pdf_path, 115.0)
    doc_out = fitz.open(out_pdf_path)
    page_out = doc_out[0]
    height_after = page_out.rect.height
    print(f"[3] Marj eklendi: {out_pdf_path}")
    print(f"    Orijinal Yükseklik: 842.0, Yeni Yükseklik: {height_after}")
    doc_out.close()

    # 4. Doğrulama Kodu Çıkarımı
    img = Image.new('RGB', (800, 200), color='white')
    d = ImageDraw.Draw(img)
    # Bizdeki font genelde arial.ttf'dir ama ImageFont.load_default() daha güvenli.
    try:
        font = ImageFont.truetype("arial.ttf", 30)
    except:
        font = ImageFont.load_default()
    d.text((50, 50), "Doğrulama Kodu: ABCD-1234-EFGH", fill='black', font=font)
    
    code = extract_verification_code(img)
    print(f"[4] Doğrulama Kodu Çıkarımı: {code}")
    print(f"    Doğru kod (ABCD-1234-EFGH) bulundu mu?: {code == 'ABCD-1234-EFGH'}")

if __name__ == "__main__":
    run_tests()
