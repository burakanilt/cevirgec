import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.ocr.engine import run_ocr
from PIL import Image

def generate_ek1_decision(data: dict, output_path: str):
    """
    EK-1 Karar Örneği şablonunu oluşturur (İlk kullanıcı yetkilendirme).
    Resmi Ticaret Bakanlığı ETDS formatına uygun şekilde tablolar içerir.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('KARAR ÖRNEĞİ', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Karar Tarihi: {data.get('karar_tarihi', '')}")
    doc.add_paragraph(f"Karar Sayısı: {data.get('karar_no', '')}")
    doc.add_paragraph(f"Toplantıya Katılanlar: {data.get('katilanlar', '')}")
    doc.add_paragraph("")
    
    doc.add_heading('1. Şirket Bilgileri', level=2)
    t_sirket = doc.add_table(rows=5, cols=2)
    t_sirket.style = 'Table Grid'
    t_sirket.rows[0].cells[0].text = "Şirket Unvanı"
    t_sirket.rows[0].cells[1].text = data.get('sirket_unvani', '')
    t_sirket.rows[1].cells[0].text = "MERSİS No"
    t_sirket.rows[1].cells[1].text = data.get('mersis_no', '')
    t_sirket.rows[2].cells[0].text = "Vergi Dairesi"
    t_sirket.rows[2].cells[1].text = data.get('vergi_dairesi', '')
    t_sirket.rows[3].cells[0].text = "Vergi Kimlik Numarası"
    t_sirket.rows[3].cells[1].text = data.get('vergi_no', '')
    t_sirket.rows[4].cells[0].text = "Bağlı Bulunduğu Ticaret Sicili Md. ve No"
    t_sirket.rows[4].cells[1].text = f"{data.get('ticaret_sicil_md', '')} - {data.get('ticaret_sicil_no', '')}"
    
    doc.add_paragraph("")
    
    doc.add_heading('2. Defter Bilgileri', level=2)
    t_defter = doc.add_table(rows=4, cols=2)
    t_defter.style = 'Table Grid'
    t_defter.rows[0].cells[0].text = "Defter Türü"
    t_defter.rows[0].cells[1].text = data.get('defter_turu', '')
    t_defter.rows[1].cells[0].text = "Hesap Dönemi"
    t_defter.rows[1].cells[1].text = data.get('hesap_donemi', '')
    t_defter.rows[2].cells[0].text = "Onay Tarihi ve Numarası"
    t_defter.rows[2].cells[1].text = f"{data.get('onay_tarihi', '')} - {data.get('onay_no', '')}"
    t_defter.rows[3].cells[0].text = "Onay Makamı"
    t_defter.rows[3].cells[1].text = data.get('onay_makami', '')
    
    doc.add_paragraph("")
    
    doc.add_heading('3. Kullanıcı Bilgileri ve Yetki Kapsamı', level=2)
    t_user = doc.add_table(rows=5, cols=2)
    t_user.style = 'Table Grid'
    t_user.rows[0].cells[0].text = "Ad Soyad"
    t_user.rows[0].cells[1].text = data.get('yetkili_ad_soyad', '')
    t_user.rows[1].cells[0].text = "TCKN"
    t_user.rows[1].cells[1].text = data.get('yetkili_tckn', '')
    t_user.rows[2].cells[0].text = "E-posta"
    t_user.rows[2].cells[1].text = data.get('yetkili_eposta', '')
    t_user.rows[3].cells[0].text = "Telefon"
    t_user.rows[3].cells[1].text = data.get('yetkili_telefon', '')
    
    # Yetki Kapsamı
    def ck(val): return "[X]" if val else "[ ]"
    yetki_str = f"{ck(data.get('yetki_kaydetme'))} Kaydetme   {ck(data.get('yetki_guncelleme'))} Güncelleme   {ck(data.get('yetki_silme'))} Silme   {ck(data.get('yetki_goruntuleme'))} Görüntüleme"
    t_user.rows[4].cells[0].text = "Yetki Kapsamı"
    t_user.rows[4].cells[1].text = yetki_str
    
    doc.add_paragraph("")
    content = doc.add_paragraph()
    content.add_run("Yukarıda bilgileri yer alan personelin, şirketimize ait belirtilen defterlerin Elektronik Ticaret Bilgi Sistemi (ETDS) üzerinde ilgili yetki kapsamlarında dijital olarak yönetilmesi için tam yetkili kılınmasına karar verilmiştir.")
    
    doc.add_paragraph("")
    sign_area = doc.add_paragraph("Şirket Yönetim Kurulu / İmzalar")
    sign_area.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(output_path)
    return output_path

def generate_ek2_user_notice(data: dict, output_path: str):
    """
    EK-2 Elektronik Defter Kullanıcısı Bildirim Formu (Sonradan yetki değişikliği)
    """
    doc = Document()
    
    title = doc.add_heading('ELEKTRONİK DEFTER KULLANICISI BİLDİRİM FORMU', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Şirket Unvanı: {data.get('sirket_unvani', '')}")
    doc.add_paragraph(f"Bildirim Tarihi: {data.get('bildirim_tarihi', '')}")
    doc.add_paragraph("")
    
    # Yeni Atanan Kullanıcılar Tablosu
    doc.add_heading('Yeni Atanan Kullanıcılar', level=2)
    new_users = data.get('yeni_kullanicilar', [])
    if new_users:
        table1 = doc.add_table(rows=1, cols=6)
        table1.style = 'Table Grid'
        hdr1 = table1.rows[0].cells
        hdr1[0].text = 'Ad Soyad'
        hdr1[1].text = 'TCKN'
        hdr1[2].text = 'E-posta'
        hdr1[3].text = 'Telefon'
        hdr1[4].text = 'Defter Türü'
        hdr1[5].text = 'Yetki Kapsamı'
        
        for u in new_users:
            r = table1.add_row().cells
            r[0].text = u.get('ad_soyad', '')
            r[1].text = u.get('tckn', '')
            r[2].text = u.get('eposta', '')
            r[3].text = u.get('telefon', '')
            r[4].text = u.get('defter_turu', '')
            r[5].text = u.get('yetki_kapsami', '')
    else:
        doc.add_paragraph("Yeni atanan kullanıcı bulunmamaktadır.")
        
    doc.add_paragraph("")
    
    # Yetkisi Kaldırılan Kullanıcılar Tablosu
    doc.add_heading('Yetkisi Kaldırılan Kullanıcılar', level=2)
    removed_users = data.get('kaldırilan_kullanicilar', [])
    if removed_users:
        table2 = doc.add_table(rows=1, cols=4)
        table2.style = 'Table Grid'
        hdr2 = table2.rows[0].cells
        hdr2[0].text = 'Ad Soyad'
        hdr2[1].text = 'TCKN'
        hdr2[2].text = 'Defter Türü'
        hdr2[3].text = 'Yetki Kapsamı'
        
        for u in removed_users:
            r = table2.add_row().cells
            r[0].text = u.get('ad_soyad', '')
            r[1].text = u.get('tckn', '')
            r[2].text = u.get('defter_turu', '')
            r[3].text = u.get('yetki_kapsami', '')
    else:
        doc.add_paragraph("Yetkisi kaldırılan kullanıcı bulunmamaktadır.")

    doc.add_paragraph("")
    
    sign_area = doc.add_paragraph("Şirket Yetkilisi\nAd Soyad / İmza")
    sign_area.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(output_path)
    return output_path

def extract_verification_code(image: Image.Image) -> str:
    """
    RapidOCR kullanarak sayfa görselinden doğrulama kodunu çıkarır.
    """
    results = run_ocr(image)
    if not results:
        return ""
        
    full_text = " ".join([text for bbox, text, conf in results])
    
    match = re.search(r'(?i)Doğrulama\s*Kodu\s*[:=\-]?\s*([A-Z0-9\-]{10,25})', full_text)
    if match:
        return match.group(1).strip()
        
    uuid_match = re.search(r'([a-fA-F0-9]{8}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{12})', full_text)
    if uuid_match:
        return uuid_match.group(1)
        
    return ""
